import csv
import json
import os
import re


class ExcelWordAI:
    def __init__(self):
        self._openpyxl = None
        self._docx = None
        self._try_imports()

    def _try_imports(self):
        try:
            import openpyxl
            self._openpyxl = openpyxl
        except ImportError:
            pass
        try:
            import docx
            self._docx = docx
        except ImportError:
            pass

    # --- Excel operations ---

    def read_excel(self, filepath, sheet_name=None):
        if self._openpyxl:
            wb = self._openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            ws = wb[sheet_name] if sheet_name else wb.active
            data = []
            for row in ws.iter_rows(values_only=True):
                data.append([str(c) if c is not None else "" for c in row])
            wb.close()
            return data
        return self._read_csv_fallback(filepath)

    def write_excel(self, filepath, data, sheet_name="Sheet1"):
        if not self._openpyxl:
            if filepath.endswith(".xlsx"):
                raise NotImplementedError("openpyxl is required for .xlsx files. Install with: pip install openpyxl")
            return self._write_csv_fallback(filepath, data)
        wb = self._openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name
        for row_idx, row in enumerate(data, 1):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        wb.save(filepath)
        wb.close()

    def analyze_excel(self, filepath, sheet_name=None):
        data = self.read_excel(filepath, sheet_name)
        if not data or len(data) < 2:
            return {"rows": len(data), "columns": len(data[0]) if data else 0}
        headers = data[0]
        col_count = len(headers)
        numeric_counts = {}
        for i, h in enumerate(headers):
            count = sum(1 for row in data[1:] if i < len(row) and _is_numeric(row[i]))
            numeric_counts[h] = count
        return {
            "rows": len(data) - 1,
            "columns": col_count,
            "headers": headers,
            "numeric_column_counts": numeric_counts,
        }

    def read_word(self, filepath):
        if not self._docx:
            if filepath.endswith(".docx"):
                raise NotImplementedError("python-docx is required. Install with: pip install python-docx")
            raise NotImplementedError("python-docx is not available")
        doc = self._docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)
        return {"paragraphs": paragraphs, "tables": tables, "paragraph_count": len(paragraphs)}

    def write_word(self, filepath, content):
        if not self._docx:
            raise NotImplementedError("python-docx is required. Install with: pip install python-docx")
        doc = self._docx.Document()
        if isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    doc.add_heading(item.get("heading", ""), level=item.get("level", 1))
                    if "text" in item:
                        doc.add_paragraph(item["text"])
                else:
                    doc.add_paragraph(str(item))
        doc.save(filepath)

    def analyze_word(self, filepath):
        content = self.read_word(filepath)
        word_count = sum(len(p.split()) for p in content["paragraphs"])
        return {
            "paragraph_count": content["paragraph_count"],
            "table_count": len(content["tables"]),
            "word_count": word_count,
        }

    def _read_csv_fallback(self, filepath):
        data = []
        with open(filepath, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            for row in reader:
                data.append(row)
        return data

    def _write_csv_fallback(self, filepath, data):
        with open(filepath, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerows(data)


def _is_numeric(val):
    try:
        float(val)
        return True
    except (ValueError, TypeError):
        return False
